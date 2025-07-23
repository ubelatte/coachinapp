# === IMPORTS ===
import streamlit as st
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from io import BytesIO
from datetime import date
import requests
import pandas as pd
import altair as alt
import gspread
from google.oauth2.service_account import Credentials
import datetime
import re

# === RESET HANDLER ===
if "reset_form" in st.session_state:
    for key in list(st.session_state.keys()):
        del st.session_state[key]
    st.rerun()

# === PAGE CONFIG ===
st.set_page_config(page_title="Mestek Coaching Generator", layout="wide")

# === PASSWORD ===
PASSWORD = "mestek"
if st.text_input("Enter password:", type="password") != PASSWORD:
    st.warning("Please enter the correct password.")
    st.stop()

# === GOOGLE SHEET SETUP ===
scope = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
service_account_info = st.secrets["gcp_service_account"]
creds = Credentials.from_service_account_info(service_account_info, scopes=scope)
gs_client = gspread.authorize(creds)
sheet = gs_client.open("Coaching Assessment Form").sheet1

# === OPENAI SETUP ===
client = OpenAI(api_key=st.secrets["openai"]["api_key"])

# === HELPER FUNCTIONS ===
def add_bold_para(doc, label, value):
    para = doc.add_paragraph()
    run = para.add_run(label)
    run.bold = True
    para.add_run(f" {value}")

def add_section_header(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)

def parse_coaching_sections(raw_text):
    sections = {}
    current_section = None
    buffer = []
    for line in raw_text.splitlines():
        line = line.strip()
        if line.endswith(":") and line[:-1] in ["Incident Summary", "Expectations Going Forward", "Tags", "Severity"]:
            if current_section and buffer:
                sections[current_section] = " ".join(buffer).strip()
                buffer = []
            current_section = line[:-1]
        elif current_section:
            buffer.append(line)
    if current_section and buffer:
        sections[current_section] = " ".join(buffer).strip()
    return sections

def add_markdown_bold_paragraph(doc, text):
    """
    Parses text with **markdown bold** syntax and adds it to the Word doc with actual bold styling.
    """
    para = doc.add_paragraph()
    bold = False
    buffer = ''
    i = 0
    while i < len(text):
        if text[i:i+2] == '**':
            if buffer:
                run = para.add_run(buffer)
                run.bold = bold
                buffer = ''
            bold = not bold
            i += 2
        else:
            buffer += text[i]
            i += 1
    if buffer:
        run = para.add_run(buffer)
        run.bold = bold

def build_coaching_doc(latest, coaching_dict):
    doc = Document()
    doc.add_heading("Employee Coaching & Counseling Form", 0)
    doc.add_paragraph(f"(Created {date.today().strftime('%m/%d/%y')})")

    doc.add_heading("Section 1 – Supervisor Entry", level=1)
    for field in [
        "Date of Incident", "Department", "Employee Name", "Supervisor Name",
        "Action to be Taken", "Issue Type", "Incident Description", "Estimated/Annual Cost",
        "Language Spoken", "Previous Coaching/Warnings", "Current Discipline Points"]:
        add_bold_para(doc, field + ":", latest.get(field, "[Missing]"))

    doc.add_page_break()
    doc.add_heading("Section 2 – Coaching Report", level=1)
    for section in ["Incident Summary", "Expectations Going Forward", "Tags", "Severity"]:
        if section in coaching_dict:
            add_section_header(doc, section + ":")
            add_markdown_bold_paragraph(doc, coaching_dict[section])  # ✅ real bold from markdown

    doc.add_paragraph("\nAcknowledgment of Receipt:")
    doc.add_paragraph(
        "I understand that this document serves as a formal record of the counseling provided. "
        "I acknowledge that the issue has been discussed with me, and I understand the expectations going forward. "
        "My signature below does not necessarily indicate agreement but confirms that I have received and reviewed this documentation.")
    doc.add_paragraph("Employee Signature: _________________________        Date: ________________")
    doc.add_paragraph("Supervisor Signature: ________________________        Date: ________________")
    return doc



def build_leadership_doc(latest, leadership_text):
    doc = Document()
    doc.add_heading("Leadership Reflection", 0)
    for field in ["Supervisor Name", "Employee Name", "Department", "Issue Type", "Date of Incident"]:
        add_bold_para(doc, field + ":", latest.get(field, "[Missing]"))

    add_section_header(doc, "AI-Generated Leadership Guidance:")

    sections = [
        "Private Reflection", "Coaching Tips", "Tone Guidance",
        "Follow-Up Recommendation", "Supervisor Accountability Tip"
    ]
    current_title = None
    buffer = []

    for line in leadership_text.splitlines() + [""]:
        stripped = line.strip()
        if stripped.endswith(":") and stripped[:-1] in sections:
            if current_title and buffer:
                doc.add_paragraph().add_run(current_title + ":").bold = True
                for para in buffer:
                    doc.add_paragraph(para)
                buffer = []
            current_title = stripped[:-1]
        elif current_title:
            buffer.append(stripped)

    if current_title and buffer:
        doc.add_paragraph().add_run(current_title + ":").bold = True
        for para in buffer:
            doc.add_paragraph(para)

    return doc


def log_submission_to_sheet(data_dict):
    timestamp = datetime.datetime.now().strftime("%m/%d/%Y %H:%M:%S")
    row = [
        timestamp,
        data_dict.get("Supervisor Name", ""),
        data_dict.get("Employee Name", ""),
        data_dict.get("Department", ""),
        data_dict.get("Date of Incident", ""),
        data_dict.get("Issue Type", ""),
        data_dict.get("Action to be Taken", ""),
        data_dict.get("Incident Description", ""),
        data_dict.get("Current Discipline Points", ""),
        data_dict.get("Estimated/Annual Cost", ""),
        data_dict.get("Language Spoken", ""),
        data_dict.get("Previous Coaching/Warnings", "")
    ]
    sheet.append_row(row, value_input_option="USER_ENTERED")



# === SESSION STATE INIT ===
if "submitted" not in st.session_state:
    st.session_state.submitted = False
    st.session_state.generated = False

# === TABS ===
tab1, tab2 = st.tabs(["📝 Coaching Form", "📊 Trend Dashboard"])




with tab1:
    with st.form("coaching_form"):
        supervisor = st.selectbox("Supervisor Name", [
            "Marty", "Nick", "Pete", "Ralph", "Steve", "Bill", "John",
            "Janitza", "Fundi", "Lisa", "Dave", "Dean"])
        employee = st.text_input("Employee Name")
        department = st.selectbox("Department", [
            "Rough In", "Paint Line (NP)", "Commercial Fabrication",
            "Baseboard Accessories", "Maintenance", "Residential Fabrication",
            "Residential Assembly/Packing", "Warehouse (55WIPR)",
            "Convector & Twin Flo", "Shipping/Receiving/Drivers",
            "Dadanco Fabrication/Assembly", "Paint Line (Dadanco)"])
        incident_date = st.date_input("Date of Incident", value=date.today())
        issue_type = st.selectbox("Issue Type", [
            "Attendance", "Safety", "Behavior", "Performance", "Policy Violation", "Recognition"])
        action_taken = st.selectbox("Action to be Taken", [
            "Coaching", "Verbal Warning", "Written Warning", "Suspension", "Termination"])
        description = st.text_area("Incident Description")
        points = st.text_input("Current Discipline Points")
        estimated_cost = st.text_input("Estimated/Annual Cost (optional)")
        language_option = st.selectbox("Language Spoken", ["English", "Spanish", "Other"])
        language = st.text_input("Please specify the language:") if language_option == "Other" else language_option
        previous = st.text_area("Previous Coaching/Warnings (if any)", placeholder="e.g., Verbal warning issued on 7/1 for tardiness.")
        submitted = st.form_submit_button("Generate Coaching Report")

    if submitted:
        st.session_state.submitted = True
        st.session_state.generated = False
        st.session_state.latest = {
            "Timestamp": date.today().isoformat(),
            "Supervisor Name": supervisor,
            "Employee Name": employee,
            "Department": department,
            "Date of Incident": incident_date.strftime("%Y-%m-%d"),
            "Issue Type": issue_type,
            "Action to be Taken": action_taken,
            "Incident Description": description,
            "Current Discipline Points": points,
            "Estimated/Annual Cost": estimated_cost,
            "Language Spoken": language,
            "Previous Coaching/Warnings": previous
        }

if st.session_state.submitted and not st.session_state.generated:
    latest = st.session_state.latest
    safe_name = latest["Employee Name"].replace(" ", "_")

    coaching_prompt = f"""
You are a workplace coaching assistant. Generate a Workplace Coaching Report using this structure and tone. Follow it exactly.

Tone & Focus Requirements:
- Be factual and objective but not cold.
- Reference Mestek policies or safety procedures if violated.
- Clearly define expected improvements.
- Use the document to educate and redirect, not punish.
- Constructive and non-punitive unless it is a formal written warning.
- Designed to correct behavior, ensure understanding, and support employee improvement.
- Avoid accusatory or vague language.
- Must be consistent with HR policies and state/federal labor laws.

Policy References:
- Counseling is a first step in the progressive discipline process (Factory Policies Packet 2025 – Performance Evaluation & Attendance).
- Documentation must include facts, dates, behaviors, and prior expectations.
- Use the Attendance and Points System for absenteeism, tardiness, or no-call/no-show (refer to Factory Policies Packet 2025 – Attendance).
- Reinforce respectful behavior for workplace conduct issues (Workplace Mutual Respect Policy).
- Reference specific Safety Procedures for PPE or Machine Guarding violations (e.g., PPE #11, Guarding #09).
- If safety is involved, cite applicable OSHA 29 CFR 1910 standards.
- Include space for employee response. Acknowledge that signatures reflect receipt, not agreement.

Structure:
Incident Summary:
On {latest['Date of Incident']}, at the {latest['Department']} location, employee {latest['Employee Name']} was involved in a situation that required supervisory intervention. The issue was identified as {latest['Issue Type']}, and the corrective action taken was {latest['Action to be Taken']}.
- Rewrite the supervisor's description in a formal, detailed, polite, and objective tone.
- Include relevant background information (e.g., point history, leaves, previous actions).
- Expand on the timeline and explain the significance of the issue and the response.
- Mention any impact to productivity or policy. 
- Include point progression or regression.
- If a cost is provided, include: "The estimated or associated cost of this issue is {latest['Estimated/Annual Cost']}."
- Conclude by noting: "Continued issues may result in progressive discipline, per Mestek guidelines."
- Do not directly copy the supervisor’s description.

Expectations Going Forward:
Clearly explain what the employee is expected to change or improve. Be firm, supportive, and specific.

Tags:
List 2-4 short keywords (e.g., attendance, policy violation, safety).

Action Taken:
Simply restate which action was taken. (e.g., coaching, verbal warning, written warning, suspension, termination, etc.)
"""




    leadership_prompt = f"""
You are a leadership coach. Write a private reflection including:
Private Reflection:
Coaching Tips:
Tone Guidance:
Follow-Up Recommendation:
Supervisor Accountability Tip:

Info:
Supervisor: {latest['Supervisor Name']}
Employee: {latest['Employee Name']}
Department: {latest['Department']}
Issue Type: {latest['Issue Type']}
Description: {latest['Incident Description']}
"""

    with st.spinner("Generating documents..."):
        coaching_response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": coaching_prompt}]
        ).choices[0].message.content.strip()

        if latest['Language Spoken'].lower() != "english":
            coaching_response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": f"Translate into {latest['Language Spoken']}\n{coaching_response}"}]
            ).choices[0].message.content.strip()

        leadership_response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": leadership_prompt}]
        ).choices[0].message.content.strip()

    coaching_sections = parse_coaching_sections(coaching_response)
    coaching_io = BytesIO()
    build_coaching_doc(latest, coaching_sections).save(coaching_io)
    coaching_io.seek(0)

    leadership_io = BytesIO()
    build_leadership_doc(latest, leadership_response).save(leadership_io)
    leadership_io.seek(0)

    try:
        log_submission_to_sheet(latest)
        st.success("✅ Submission logged to Google Sheet.")
    except Exception as e:
        st.error(f"❌ Could not log to Google Sheet.\n{e}")

    st.session_state.coaching_io = coaching_io
    st.session_state.leadership_io = leadership_io
    st.session_state.safe_name = safe_name
    st.session_state.generated = True

if st.session_state.get("generated", False):
    col1, col2 = st.columns(2)
    with col1:
        st.download_button("📄 Download Coaching Doc", data=st.session_state.coaching_io,
                           file_name=f"{st.session_state.safe_name}_coaching.docx")
    with col2:
        st.download_button("📄 Download Leadership Doc", data=st.session_state.leadership_io,
                           file_name=f"{st.session_state.safe_name}_leadership.docx")


# === TREND DASHBOARD ===
with tab2:
    st.header("📊 Coaching Trend Dashboard")
    try:
        df = pd.DataFrame(sheet.get_all_records())
        df["Date of Incident"] = pd.to_datetime(df["Date of Incident"], errors="coerce")

        min_date = df["Date of Incident"].min()
        max_date = df["Date of Incident"].max()
        start_date, end_date = st.date_input("Filter by Date Range", [min_date, max_date], key="date_range_filter")

        if start_date and end_date:
            df = df[(df["Date of Incident"] >= pd.to_datetime(start_date)) & (df["Date of Incident"] <= pd.to_datetime(end_date))]

        filter_action = st.selectbox(
            "Filter by Action Taken",
            ["All"] + df["Action to be Taken"].dropna().unique().tolist(),
            key="trend_action_filter"
        )
        if filter_action != "All":
            df = df[df["Action to be Taken"] == filter_action]

        st.dataframe(df)

        st.subheader("Issue Type Count")
        issue_counts = df["Issue Type"].value_counts().reset_index()
        issue_counts.columns = ["Issue Type", "Count"]
        chart = alt.Chart(issue_counts).mark_bar().encode(
            x=alt.X("Issue Type", sort="-y"),
            y="Count",
            tooltip=["Issue Type", "Count"]
        ).properties(width=600, height=400)
        st.altair_chart(chart, use_container_width=True)

        st.subheader("Actions Over Time")
        df["Date Only"] = df["Date of Incident"].dt.date
        trend = df.groupby(["Date Only", "Action to be Taken"]).size().unstack(fill_value=0)
        st.line_chart(trend)

    except Exception as e:
        st.error(f"❌ No Info Logged: {e}")
